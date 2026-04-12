"""Parse Excel and DocX supplier files, extract text, then use LLM to structure data."""
import io
import openpyxl
import docx
from pathlib import Path
from .llm_tasks import extract_items_from_text


def _read_xlsx(path: str) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Лист: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _read_docx(path: str) -> str:
    doc = docx.Document(path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


async def parse_supplier_file(file_path: str) -> list[dict]:
    """
    Parse supplier file (xlsx/xls/docx) and return structured item list via LLM.
    Each item: {name, registry_number, okpd2_code, quantity, characteristics: [{name, value}]}
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext in (".xlsx", ".xls"):
        raw_text = _read_xlsx(file_path)
    elif ext == ".docx":
        raw_text = _read_docx(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    items = await extract_items_from_text(raw_text)
    return items
